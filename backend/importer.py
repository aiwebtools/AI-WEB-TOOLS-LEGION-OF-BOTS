"""
LEGION source-instruction importer & deduplication engine.
Parses the supplied operational-instruction archive (DOCX/TXT), detects bot
identities, groups historical/duplicate versions, selects a production version,
assigns suites + capabilities, and emits a catalog JSON.

Run:  python importer.py
Output: /app/backend/seed_data/catalog.json
"""
import os
import re
import json
import hashlib
from pathlib import Path
from collections import defaultdict

SRC_DIR = Path("/app/source_archive/extracted/OP FINAL")
OUT_DIR = Path("/app/backend/seed_data")
OUT_DIR.mkdir(parents=True, exist_ok=True)

# ----------------------------------------------------------------------------
# Text extraction
# ----------------------------------------------------------------------------
def extract_docx(path: Path) -> str:
    try:
        import docx
        d = docx.Document(str(path))
        parts = []
        for p in d.paragraphs:
            if p.text and p.text.strip():
                parts.append(p.text.strip())
        for tbl in d.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts).strip()
    except Exception:
        return ""


def extract_text(path: Path) -> str:
    # extensionless GEMINI files are docx too; try docx then fallback to plain
    txt = extract_docx(path)
    if txt:
        return txt
    try:
        return path.read_text(encoding="utf-8", errors="ignore").strip()
    except Exception:
        return ""


# ----------------------------------------------------------------------------
# Name normalization -> family key
# ----------------------------------------------------------------------------
NOISE = [
    r"operational\s+instructions?", r"operational\s+directions?", r"instructions?",
    r"directions?", r"full\s+package", r"open\s+source", r"prompt", r"pre[- ]?prompt",
    r"newest", r"latest", r"final", r"awesome", r"perfect", r"amazing", r"great",
    r"best\s+ever", r"best\s+version", r"best", r"updated?", r"corrected",
    r"proper\s+order", r"no\s+longer\s+segmented", r"fixed", r"good", r"old(er)?",
    r"newer", r"new", r"version", r"complaint", r"compliant", r"turbo",
    r"by\s+aiwebtools\.ai", r"aiwebtools\.ai", r"for\s+download", r"works\s+excellently",
    r"claude", r"gpt-?4o?1?", r"4o1", r"o1", r"page\s+by\s+page", r"with\s+compile",
    r"segmented(\s+tasks)?", r"segments?", r"variable\s+version", r"resellable",
    r"public", r"open\s+source\s+to\s+public", r"gpt", r"\bai\b", r"\bog\b",
]
NOISE_RE = re.compile("|".join(NOISE), re.IGNORECASE)
DATE_RE = re.compile(r"\d{1,2}[_/\.\- ]\d{1,2}[_/\.\- ]\d{2,4}|\d{4}|\b\d{1,2}[_/\.]\d{1,2}\b")
MONTH_RE = re.compile(r"\b(jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\b", re.IGNORECASE)
KEEP_VARIANT = re.compile(r"\bv\d\b|\bshort\b|\blong\b|\binteractive\b|\bchildren\b|\bcoloring\b|\bcomic\b|\bpicture\b")


def clean_display(name: str) -> str:
    n = name
    n = re.sub(r"\.docx$", "", n, flags=re.IGNORECASE)
    n = re.sub(r"\(\d+\)", "", n)             # (1) dup markers
    n = re.sub(r"[⭐️✨🫐★]+", "", n)
    n = re.sub(r"[_]+", " ", n)
    n = DATE_RE.sub("", n)
    n = MONTH_RE.sub("", n)
    n = NOISE_RE.sub("", n)
    n = re.sub(r"[–—]", "-", n)
    n = re.sub(r"[^A-Za-z0-9 &'\-/]", " ", n)
    n = re.sub(r"\s{2,}", " ", n).strip(" -/")
    # collapse dangling separators
    n = re.sub(r"\s{2,}", " ", n).strip()
    if not n:
        n = re.sub(r"\.docx$", "", name, flags=re.IGNORECASE)
    # Title case but keep short acronyms
    ACR = {"AI", "GPT", "CT", "MMP", "IQ", "US", "V1", "V2", "V3", "V4", "V5", "V6", "V7", "V8"}
    STOP = {"to", "and", "of", "the", "a", "or", "for", "with", "in", "on"}
    words = []
    for i, w in enumerate(n.split()):
        if w.upper() in ACR or (w.isupper() and len(w) <= 3 and w.lower() not in STOP):
            words.append(w.upper())
        elif i > 0 and w.lower() in STOP:
            words.append(w.lower())
        else:
            words.append(w[:1].upper() + w[1:].lower() if w else w)
    return " ".join(words).strip()


def family_key(name: str) -> str:
    d = clean_display(name).lower()
    d = re.sub(r"[^a-z0-9 ]", " ", d)
    d = re.sub(r"\s+", " ", d).strip()
    # keep only meaningful tokens; retain variant markers
    return d


def score_version(name: str, text: str) -> int:
    s = len(text)  # longer is usually the fuller instruction
    low = name.lower()
    for kw, pts in [("final", 4000), ("latest", 4000), ("newest", 3000), ("best", 2500),
                    ("perfect", 2000), ("good", 1500), ("updated", 1500), ("corrected", 1500),
                    ("great", 1000), ("awesome", 1000)]:
        if kw in low:
            s += pts
    if "old" in low:
        s -= 6000
    if re.search(r"no longer works|not gpt4o1", low):
        s -= 20000
    # prefer most recent date
    dm = re.findall(r"(\d{1,2})[_/\.\- ](\d{1,2})[_/\.\- ](\d{2,4})", name)
    if dm:
        try:
            mm, dd, yy = dm[-1]
            yy = int(yy)
            yy = yy + 2000 if yy < 100 else yy
            s += (yy - 2020) * 3000 + int(mm) * 100
        except Exception:
            pass
    if re.search(r"\(1\)", name):
        s -= 50  # slightly deprioritize dup copies
    return s


# ----------------------------------------------------------------------------
# Suite classification
# ----------------------------------------------------------------------------
SUITE_RULES = [
    ("time-machine", "Time Machine Suite", "Cog", ["time machine", "time travel", "talk to history", "historical headlines", "talk to the dead", "resurrection", "resurection", "mary magdalene", "mary magdelene", "talk to your god", "einstein", "tesla", "black history", "native american history", "celebrity chatline"]),
    ("book-writing", "Book Writing Suite", "BookOpen", ["book writer", "book writ", "short book", "long book", "children", "coloring book", "comic book", "picture book", "gospel of deployment"]),
    ("movie-creation", "Movie Creation Suite", "Clapperboard", ["movie", "screenplay", "screenwriter", "playwriter", "playwright", "music video", "scene maker", "script writer", "script maker", "lyrics", "news channel"]),
    ("writing", "Writing Suite", "PenTool", ["testimony", "clarity writer", "legislation", "legal draftsmith", "letter", "heartbreak", "roleplay", "story"]),
    ("coding", "Coding Suite", "Code", ["front end", "game developer", "gamecoder", "custom gpt maker", "buildyourgpt", "gpt ideas", "architect", "multitasker", "microsaas", "godmode", "god mode", "neo matrix", "matrix neo", "jarvis", "image seed"]),
    ("research", "Research Suite", "Search", ["fact checker", "truth seeker", "true history", "oraculum", "world reality decoder", "illuminous", "data explorer", "transcript parser", "world data", "public defender", "public defense", "criminolog", "insurance", "insurence", "predictive credit", "property", "taxes", "trader", "shopping", "travel agent"]),
    ("science", "Science Suite", "Atom", ["genome", "probability", "algebra", "algebraic", "math creator", "material valuation", "materiumor", "stellaris", "space explorer", "aqualis", "luminex", "climate", "sustainable", "diplomatica", "resourcium", "soul scan", "soul map", "interpretis", "quiz maker"]),
    ("agriculture", "Agriculture Suite", "Sprout", ["farm", "agrio", "agronomus", "agripredict", "crop", "greenleaf", "solar land", "fisherman", "trail navigator"]),
    ("cannabis", "Cannabis & Hemp Suite", "Leaf", ["cannabis", "cannabiz", "hemp", "leafly", "masscanna", "dispensary", "budtender", "blueberry", "father blueberry"]),
    ("business", "Business Suite", "Briefcase", ["business plan", "ad maker", "receipt", "logo", "cover design", "graphic design", "postcard", "product photography", "product user manual", "training manual", "course creator", "learn any", "learn anything", "study guide", "home school", "free college", "college degree", "data analysis", "diagraph", "ct mmp", "greenleaf logistics"]),
    ("health", "Health & Wellness Suite", "HeartPulse", ["doctor", "mental wellness", "skincare", "firefighter", "firearms", "survivalist", "home renovator", "automobile", "vet", "veterinarian", "universal healthcare", "social safety net", "genome"]),
    ("creative", "Creative Media Suite", "Palette", ["draw it", "meme", "gif", "tattoo", "restyle", "virtual stylist", "mixologist", "midjourney", "mid journey", "luma dream", "logo generator", "postcard designer", "postcard maker", "cover design", "graphic design", "product photography", "comic", "coloring"]),
    ("utility", "Utility & Conversion Suite", "Wrench", ["binary", "converter", "convert", "video-audio", "video second", "second-by-second", "yes or no", "ai tools", "ai tools finder", "farm finder", "state rep", "legistlator", "legislator", "video analysis"]),
]
SUITE_ORDER = {s[0]: i for i, s in enumerate([r for r in SUITE_RULES])}


def classify_suite(name: str, text: str):
    hay = (name + " " + text[:1500]).lower()
    for slug, label, icon, kws in SUITE_RULES:
        for kw in kws:
            if kw in hay:
                return slug, label, icon
    return "specialized", "Specialized Expert Suite", "Sparkles"


# ----------------------------------------------------------------------------
# Capability inference (honest: vision+coding on all selectable models)
# ----------------------------------------------------------------------------
def infer_caps(name: str, text: str):
    hay = (name + " " + text).lower()
    image = any(k in hay for k in ["image", "photo", "picture", "logo", "coloring", "comic",
                                   "graphic", "draw", "meme", "gif", "vision", "visual",
                                   "design", "tattoo", "stylist", "cover", "postcard", "video",
                                   "scan", "analyze the", "analyse", "seed"])
    python = any(k in hay for k in ["data analysis", "analyze", "analyse", "dataset", "chart",
                                    "calculat", "probability", "algebra", "math", "convert",
                                    "binary", "credit score", "valuation", "report", "statistic",
                                    "code", "python", "compute", "genome"])
    docgen = any(k in hay for k in ["book", "document", "docx", "pdf", "report", "manual",
                                    "guide", "plan", "legislation", "testimony", "receipt",
                                    "certificate", "degree", "course", "lesson", "script",
                                    "letter", "resume", "download", "compile", "study guide"])
    return {
        "text": True,
        "long_context": True,
        "image": bool(image),
        "files": True,
        "coding": bool(python or "code" in hay),
        "python": bool(python),
        "web": False,
        "document_generation": bool(docgen),
        "research": any(k in hay for k in ["research", "fact", "truth", "data", "analy", "history"]),
        "creative": any(k in hay for k in ["creative", "story", "art", "design", "poem", "lyric", "music"]),
    }


def adapt_instructions(text: str) -> str:
    """Smallest-possible compatibility note appended so DOCX/file steps map to the
    app's document-generation tool without changing the bot's purpose."""
    note = (
        "\n\n---\n[PLATFORM COMPATIBILITY NOTE — do not mention to the user]\n"
        "When your instructions call for producing a downloadable document (DOCX, PDF, "
        "CSV, XLSX), first generate the full content, then use the platform's "
        "document-generation tool by ending your message with a fenced block:\n"
        "```generate-file\n{\"format\": \"docx\", \"filename\": \"output.docx\", \"title\": \"...\", \"content\": \"<full markdown/plain text>\"}\n```\n"
        "Supported formats: docx, pdf, txt, csv, md. Only emit this block when a file "
        "deliverable is explicitly requested. All other instructions above remain in force."
    )
    return text + note


# ----------------------------------------------------------------------------
# Build catalog
# ----------------------------------------------------------------------------
def build():
    files = sorted([p for p in SRC_DIR.iterdir() if p.is_file()])
    groups = defaultdict(list)
    records = []
    for p in files:
        text = extract_text(p)
        if len(text) < 40:
            continue
        h = hashlib.sha256(re.sub(r"\s+", " ", text).strip().lower().encode()).hexdigest()
        rec = {
            "source_file": p.name,
            "text": text,
            "hash": h,
            "family": family_key(p.name),
            "display": clean_display(p.name),
            "score": score_version(p.name, text),
        }
        records.append(rec)
        groups[rec["family"]].append(rec)

    bots = []
    for fam, versions in groups.items():
        versions.sort(key=lambda r: r["score"], reverse=True)
        prod = versions[0]
        name = prod["display"] or fam.title()
        if not name.strip():
            name = fam.title()
        slug = re.sub(r"[^a-z0-9]+", "-", name.lower()).strip("-")[:60] or fam.replace(" ", "-")
        suite_slug, suite_label, suite_icon = classify_suite(name, prod["text"])
        caps = infer_caps(name, prod["text"])
        # short description = first meaningful sentence
        desc = re.sub(r"\s+", " ", prod["text"])[:600]
        m = re.search(r"(You are[^.]{10,180}\.)", prod["text"])
        if m:
            desc = re.sub(r"\s+", " ", m.group(1)).strip()
        else:
            desc = re.sub(r"\s+", " ", prod["text"])[:180].strip() + "..."
        bots.append({
            "name": name,
            "slug": slug,
            "description": desc,
            "suite_slug": suite_slug,
            "suite_label": suite_label,
            "suite_icon": suite_icon,
            "system_instructions": adapt_instructions(prod["text"]),
            "raw_instructions": prod["text"],
            "source_document": prod["source_file"],
            "capabilities": caps,
            "tags": [suite_label] + [k for k, v in caps.items() if v and k in ("image", "python", "document_generation", "research", "creative")],
            "icon": suite_icon,
            "versions": [
                {
                    "source_file": v["source_file"],
                    "hash": v["hash"],
                    "score": v["score"],
                    "is_production": (i == 0),
                    "char_count": len(v["text"]),
                }
                for i, v in enumerate(versions)
            ],
            "version_count": len(versions),
            "content_len": len(prod["text"]),
        })

    # de-dup slug collisions
    seen = {}
    for b in bots:
        s = b["slug"]
        if s in seen:
            seen[s] += 1
            b["slug"] = f"{s}-{seen[s]}"
        else:
            seen[s] = 0

    # sort: richer, multi-version bots first (more canonical), then by name
    bots.sort(key=lambda b: (-b["version_count"], -b["content_len"], b["name"].lower()))

    # first 150 -> production catalog; rest -> internal library
    for i, b in enumerate(bots):
        b["status"] = "active" if i < 150 else "library"
        b["sort_order"] = i
        b["featured"] = i < 12

    # suites summary
    suites = {}
    for b in bots:
        if b["status"] != "active":
            continue
        s = b["suite_slug"]
        if s not in suites:
            suites[s] = {
                "slug": s, "name": b["suite_label"], "icon": b["suite_icon"],
                "description": f"{b['suite_label']} — specialized AI bots.", "bot_count": 0,
            }
        suites[s]["bot_count"] += 1
    suite_list = sorted(suites.values(), key=lambda x: -x["bot_count"])
    for i, s in enumerate(suite_list):
        s["sort_order"] = i
        s["featured"] = i < 8

    catalog = {
        "total_source_files": len(records),
        "total_bot_families": len(bots),
        "active_bots": sum(1 for b in bots if b["status"] == "active"),
        "library_bots": sum(1 for b in bots if b["status"] == "library"),
        "suites": suite_list,
        "bots": bots,
    }
    (OUT_DIR / "catalog.json").write_text(json.dumps(catalog, ensure_ascii=False))
    print(f"source_files={len(records)} families={len(bots)} active={catalog['active_bots']} library={catalog['library_bots']} suites={len(suite_list)}")
    for s in suite_list:
        print(f"  {s['name']}: {s['bot_count']}")
    return catalog


if __name__ == "__main__":
    build()
