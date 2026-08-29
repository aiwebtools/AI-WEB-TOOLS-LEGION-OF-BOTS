"""Tools: document generation (docx/pdf/txt/csv/md) and sandboxed Python execution."""
import io
import os
import re
import csv
import json
import uuid
import base64
import subprocess
import tempfile
from datetime import datetime, timezone

GEN_DIR = "/app/backend/generated_files"
os.makedirs(GEN_DIR, exist_ok=True)

GEN_BLOCK_RE = re.compile(r"```generate-file\s*(\{.*?\})\s*```", re.DOTALL)


def extract_file_request(text: str):
    """Find a generate-file block in the assistant output; return (spec, cleaned_text)."""
    m = GEN_BLOCK_RE.search(text)
    if not m:
        return None, text
    try:
        spec = json.loads(m.group(1))
    except Exception:
        return None, text
    cleaned = GEN_BLOCK_RE.sub("", text).strip()
    return spec, cleaned


def _make_docx(title, content) -> bytes:
    import docx
    d = docx.Document()
    if title:
        d.add_heading(title, level=0)
    for line in content.split("\n"):
        s = line.rstrip()
        if s.startswith("### "):
            d.add_heading(s[4:], level=3)
        elif s.startswith("## "):
            d.add_heading(s[3:], level=2)
        elif s.startswith("# "):
            d.add_heading(s[2:], level=1)
        elif s.startswith(("- ", "* ")):
            d.add_paragraph(s[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", s):
            d.add_paragraph(re.sub(r"^\d+\.\s", "", s), style="List Number")
        else:
            d.add_paragraph(s)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _make_pdf(title, content) -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.units import inch
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, topMargin=inch, bottomMargin=inch)
    styles = getSampleStyleSheet()
    flow = []
    if title:
        flow.append(Paragraph(title, styles["Title"]))
        flow.append(Spacer(1, 12))
    for line in content.split("\n"):
        s = line.strip()
        if not s:
            flow.append(Spacer(1, 8))
            continue
        s = s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if s.startswith("# "):
            flow.append(Paragraph(s[2:], styles["Heading1"]))
        elif s.startswith("## "):
            flow.append(Paragraph(s[3:], styles["Heading2"]))
        elif s.startswith("### "):
            flow.append(Paragraph(s[4:], styles["Heading3"]))
        else:
            flow.append(Paragraph(s, styles["BodyText"]))
    doc.build(flow)
    return buf.getvalue()


def _make_csv(content) -> bytes:
    buf = io.StringIO()
    writer = csv.writer(buf)
    for line in content.strip().split("\n"):
        writer.writerow([c.strip() for c in re.split(r"\t|,|\|", line)])
    return buf.getvalue().encode("utf-8")


def generate_file(spec: dict) -> dict:
    fmt = (spec.get("format") or "txt").lower()
    title = spec.get("title", "")
    content = spec.get("content", "")
    filename = spec.get("filename") or f"output.{fmt}"
    if not filename.lower().endswith("." + fmt):
        filename = f"{os.path.splitext(filename)[0]}.{fmt}"

    if fmt == "docx":
        data = _make_docx(title, content)
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif fmt == "pdf":
        data = _make_pdf(title, content)
        mime = "application/pdf"
    elif fmt == "csv":
        data = _make_csv(content)
        mime = "text/csv"
    elif fmt in ("md", "markdown"):
        data = ((f"# {title}\n\n" if title else "") + content).encode("utf-8")
        mime = "text/markdown"
        fmt = "md"
    else:
        data = ((f"{title}\n\n" if title else "") + content).encode("utf-8")
        mime = "text/plain"
        fmt = "txt"

    file_id = str(uuid.uuid4())
    disk = os.path.join(GEN_DIR, f"{file_id}.{fmt}")
    with open(disk, "wb") as f:
        f.write(data)
    return {
        "id": file_id,
        "filename": filename,
        "format": fmt,
        "mime": mime,
        "size": len(data),
        "disk_path": disk,
        "created_at": datetime.now(timezone.utc).isoformat(),
    }


PY_HARNESS = """
import sys, io, contextlib, resource
resource.setrlimit(resource.RLIMIT_AS, (512*1024*1024, 512*1024*1024))
resource.setrlimit(resource.RLIMIT_CPU, (8, 8))
_out = io.StringIO()
try:
    with contextlib.redirect_stdout(_out), contextlib.redirect_stderr(_out):
        exec(compile(open(sys.argv[1]).read(), 'user_code.py', 'exec'), {'__name__': '__main__'})
except Exception as e:
    import traceback
    _out.write('\\n' + traceback.format_exc())
sys.stdout.write(_out.getvalue())
"""


def run_python(code: str, timeout: int = 12) -> dict:
    """Execute Python in an isolated subprocess with resource + time limits."""
    with tempfile.TemporaryDirectory() as tmp:
        code_path = os.path.join(tmp, "code.py")
        harness_path = os.path.join(tmp, "harness.py")
        with open(code_path, "w") as f:
            f.write(code)
        with open(harness_path, "w") as f:
            f.write(PY_HARNESS)
        try:
            proc = subprocess.run(
                ["python", harness_path, code_path],
                capture_output=True, text=True, timeout=timeout, cwd=tmp,
                env={"PATH": os.environ.get("PATH", ""), "HOME": tmp},
            )
            output = proc.stdout or proc.stderr
            return {"success": proc.returncode == 0, "output": output[:20000] or "(no output)"}
        except subprocess.TimeoutExpired:
            return {"success": False, "output": f"Execution timed out after {timeout}s."}
        except Exception as e:
            return {"success": False, "output": f"Execution error: {e}"}
